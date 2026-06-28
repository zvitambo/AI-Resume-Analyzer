import io
import pdfplumber
from docx import Document
import io
from typing import Optional, Callable
from app.core.exceptions import TexExtracionError



def extract_text_from_pdf(file_bytes) -> str:
    """
    Extract text from PDF bytes using pdfplumber for better formatting.
    Requires: pip install pdfplumber
    """

    try:
        pdf_file = io.BytesIO(file_bytes)
        
        with pdfplumber.open(pdf_file) as pdf:
            text = ""
            for page in pdf.pages:
                text += page.extract_text() or "" + "\n"
        
        return text.strip()
    except Exception as e:
       raise TexExtracionError("Failed to extract text from PDF") from e


def extract_text_from_docx(
    file_bytes: bytes,
    progress_callback: Optional[Callable[[int, int], None]] = None
) -> str:
    """
    Extract text from DOCX with progress tracking.
    
    Args:
        file_bytes: DOCX file content as bytes
        progress_callback: Function called with (current, total) during processing
        
    Returns:
        Extracted text as a single string
    """
    if not file_bytes:
        raise ValueError("No file bytes provided")
    
    try:
        doc_stream = io.BytesIO(file_bytes)
        doc = Document(doc_stream)
        
        # Count total elements (paragraphs + table cells)
        total_elements = len(doc.paragraphs)
        for table in doc.tables:
            total_elements += len(table.rows) * len(table.rows[0].cells) if table.rows else 0
        
        text_parts = []
        processed = 0
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
            processed += 1
            if progress_callback:
                progress_callback(processed, total_elements)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_data.append(cell_text)
                    processed += 1
                    if progress_callback:
                        progress_callback(processed, total_elements)
                if row_data:
                    text_parts.append(" | ".join(row_data))
        
        return "\n".join(text_parts)
        
    except Exception as e:
        raise TexExtracionError("Failed to extract text from DOCX") from e
    